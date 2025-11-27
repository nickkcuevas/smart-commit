#!/usr/bin/env python3
"""
Script para generar el documento Word (.docx) de la guía de Smart Commit
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import sys


def add_hyperlink(paragraph, text, url):
    """Añade un hipervínculo a un párrafo"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Estilo del hipervínculo
    c = OxmlElement('w:color')
    c.set(qn('w:val'), "0563C1")
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink


def set_run_font(run, font_name='Calibri', font_size=11, bold=False, color=None):
    """Configura la fuente de un run"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def create_document():
    """Crea el documento Word completo"""
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # TÍTULO PRINCIPAL
    title = doc.add_heading('Smart Commit - Guía Completa para Usuarios No Técnicos', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del documento
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(f'Versión: 1.0 | Fecha: {datetime.now().strftime("%Y-%m-%d")} | Autor: Equipo de Desarrollo')
    set_run_font(info_run, font_size=9, color=RGBColor(128, 128, 128))
    
    doc.add_paragraph()  # Espacio
    
    # TABLA DE CONTENIDOS
    doc.add_heading('Tabla de Contenidos', 1)
    toc_items = [
        '¿Qué es Smart Commit?',
        'Información General del Proyecto',
        'Estadísticas Principales',
        'Cambios Significativos',
        'Advertencias y Problemas Potenciales',
        'Archivos por Tipo',
        'Detalles por Archivo',
        'Crear un Commit',
        'Ejemplos Prácticos',
        'Preguntas Frecuentes',
        'Glosario de Términos'
    ]
    
    for i, item in enumerate(toc_items, 1):
        p = doc.add_paragraph(f'{i}. {item}', style='List Number')
        p.style.font.size = Pt(11)
    
    doc.add_page_break()
    
    # SECCIÓN 1: ¿Qué es Smart Commit?
    doc.add_heading('¿Qué es Smart Commit?', 1)
    
    p = doc.add_paragraph()
    p.add_run('Smart Commit es una herramienta visual que te muestra ').bold = False
    p.add_run('exactamente qué cambios').bold = True
    p.add_run(' estás a punto de guardar en el código antes de hacerlo.')
    
    doc.add_paragraph('Piensa en ella como un "resumen ejecutivo" que te dice:', style='List Bullet')
    doc.add_paragraph('Qué archivos modificaste', style='List Bullet 2')
    doc.add_paragraph('Qué tan grandes son los cambios', style='List Bullet 2')
    doc.add_paragraph('Si hay algo importante que debes revisar', style='List Bullet 2')
    doc.add_paragraph('Si hay posibles problemas o riesgos', style='List Bullet 2')
    
    doc.add_heading('¿Por qué es útil?', 2)
    doc.add_paragraph('Te ayuda a entender el alcance de tus cambios', style='List Bullet')
    doc.add_paragraph('Te alerta sobre posibles problemas antes de guardar', style='List Bullet')
    doc.add_paragraph('Te da confianza de que estás guardando lo correcto', style='List Bullet')
    doc.add_paragraph('Facilita la revisión de código', style='List Bullet')
    
    # SECCIÓN 2: Información General
    doc.add_heading('Información General del Proyecto', 1)
    
    doc.add_heading('Branch (Rama)', 2)
    doc.add_paragraph('El nombre de la versión del código en la que estás trabajando actualmente.')
    
    doc.add_heading('Ejemplos comunes:', 3)
    examples = [
        ('dev', 'Versión de desarrollo'),
        ('main o master', 'Versión principal del proyecto'),
        ('feature/nueva-funcionalidad', 'Nueva funcionalidad en desarrollo'),
        ('bugfix/correccion-error', 'Corrección de un error')
    ]
    
    for name, desc in examples:
        p = doc.add_paragraph()
        p.add_run(f'{name} - ').bold = True
        p.add_run(desc)
    
    doc.add_paragraph('Es como trabajar en una carpeta separada. Tus cambios no afectan el código principal hasta que los mezcles. Esto permite que varias personas trabajen en el mismo proyecto sin conflictos.')
    
    doc.add_heading('Estado del Proyecto', 2)
    
    p = doc.add_paragraph()
    p.add_run('"Listo para commit"').bold = True
    doc.add_paragraph('Tienes archivos preparados y listos para guardar. Puedes proceder a crear el commit.', style='List Bullet 2')
    
    p = doc.add_paragraph()
    p.add_run('"No hay archivos staged"').bold = True
    doc.add_paragraph('No hay archivos seleccionados para guardar. Necesitas primero seleccionar qué archivos quieres incluir. Puedes usar el botón "Stage All" para añadir todos los archivos modificados.', style='List Bullet 2')
    
    # SECCIÓN 3: Estadísticas Principales
    doc.add_heading('Estadísticas Principales', 1)
    doc.add_paragraph('Estas cuatro tarjetas te dan un vistazo rápido del tamaño y alcance de tus cambios.')
    
    doc.add_heading('Tarjeta 1: Archivos (Color Azul)', 2)
    doc.add_paragraph('Muestra el número total de archivos diferentes que modificaste.')
    
    doc.add_heading('Ejemplos:', 3)
    doc.add_paragraph('3 archivos - Cambio pequeño y enfocado', style='List Bullet')
    doc.add_paragraph('15 archivos - Cambio de tamaño medio', style='List Bullet')
    doc.add_paragraph('50+ archivos - Cambio grande que afecta múltiples áreas', style='List Bullet')
    
    doc.add_heading('Cómo interpretarlo:', 3)
    doc.add_paragraph('1-5 archivos: Cambio pequeño, fácil de revisar', style='List Bullet')
    doc.add_paragraph('6-20 archivos: Cambio medio, requiere revisión cuidadosa', style='List Bullet')
    doc.add_paragraph('21+ archivos: Cambio grande, considera dividirlo en partes más pequeñas', style='List Bullet')
    
    doc.add_heading('Tarjeta 2: Añadidas (Color Verde)', 2)
    doc.add_paragraph('Muestra el número total de líneas de código NUEVAS que agregaste. Siempre se muestra con un signo +.')
    
    doc.add_heading('Ejemplos:', 3)
    doc.add_paragraph('+50 líneas - Agregaste 50 líneas nuevas de código', style='List Bullet')
    doc.add_paragraph('+250 líneas - Agregaste una cantidad significativa de código nuevo', style='List Bullet')
    doc.add_paragraph('+1000 líneas - Agregaste una gran cantidad de código (puede ser una nueva funcionalidad importante)', style='List Bullet')
    
    doc.add_heading('Cómo interpretarlo:', 3)
    doc.add_paragraph('+1 a +100 líneas: Cambio pequeño a mediano', style='List Bullet')
    doc.add_paragraph('+101 a +500 líneas: Cambio significativo, probablemente nueva funcionalidad', style='List Bullet')
    doc.add_paragraph('+501+ líneas: Cambio muy grande, asegúrate de que sea necesario', style='List Bullet')
    
    doc.add_paragraph('Indica cuánto código nuevo escribiste. Muchas líneas añadidas pueden significar nueva funcionalidad, nuevas características, o mejoras importantes.')
    
    doc.add_heading('Tarjeta 3: Eliminadas (Color Rojo)', 2)
    doc.add_paragraph('Muestra el número total de líneas de código que ELIMINASTE. Siempre se muestra con un signo -.')
    
    doc.add_heading('Cómo interpretarlo:', 3)
    doc.add_paragraph('-1 a -50 líneas: Eliminación pequeña, probablemente limpieza', style='List Bullet')
    doc.add_paragraph('-51 a -200 líneas: Eliminación significativa, puede ser refactorización', style='List Bullet')
    doc.add_paragraph('-201+ líneas: Eliminación grande, asegúrate de que no eliminaste funcionalidad importante', style='List Bullet')
    
    doc.add_paragraph('Muestra qué código removiste. Esto puede ser bueno (código viejo que ya no se necesita, código duplicado, limpieza) o requiere revisión (asegúrate de que no eliminaste funcionalidad que aún se necesita).')
    
    doc.add_heading('Tarjeta 4: Cambio Neto (Color Verde o Rojo)', 2)
    doc.add_paragraph('Muestra la diferencia entre lo que añadiste y lo que eliminaste. Puede ser positivo (verde) o negativo (rojo).')
    
    doc.add_paragraph('Cómo se calcula: Cambio Neto = Líneas Añadidas - Líneas Eliminadas')
    
    doc.add_heading('Ejemplos:', 3)
    doc.add_paragraph('+130 líneas - Añadiste 250, eliminaste 120 → Neto: +130 (el proyecto creció)', style='List Bullet')
    doc.add_paragraph('-50 líneas - Añadiste 100, eliminaste 150 → Neto: -50 (el proyecto se redujo)', style='List Bullet')
    doc.add_paragraph('+0 líneas - Añadiste 200, eliminaste 200 → Neto: 0 (mismo tamaño)', style='List Bullet')
    
    doc.add_heading('Cómo interpretarlo:', 3)
    
    p = doc.add_paragraph()
    p.add_run('Número Positivo (Verde):').bold = True
    doc.add_paragraph('El proyecto creció en tamaño. Agregaste más código del que quitaste. Normal cuando agregas nueva funcionalidad.', style='List Bullet 2')
    
    p = doc.add_paragraph()
    p.add_run('Número Negativo (Rojo):').bold = True
    doc.add_paragraph('El proyecto se redujo en tamaño. Eliminaste más código del que agregaste. Puede ser bueno (limpieza, refactorización) o malo (eliminaste funcionalidad por error).', style='List Bullet 2')
    
    # SECCIÓN 4: Cambios Significativos
    doc.add_heading('Cambios Significativos', 1)
    doc.add_paragraph('Esta sección te alerta sobre cambios importantes que podrían tener un impacto significativo en el sistema.')
    
    doc.add_heading('🔄 Migración de Base de Datos', 2)
    doc.add_paragraph('Cambios en cómo se guarda y organiza la información en la base de datos.')
    
    p = doc.add_paragraph('Ejemplo de mensaje: ')
    code_run = p.add_run('🔄 Migración de base de datos: zdash/revi/orders/migrations/0023_add_user_field.py')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    
    doc.add_heading('¿Por qué es importante?', 3)
    doc.add_paragraph('Afecta cómo se almacenan los datos', style='List Bullet')
    doc.add_paragraph('Puede requerir pasos especiales al desplegar', style='List Bullet')
    doc.add_paragraph('Puede afectar datos existentes', style='List Bullet')
    doc.add_paragraph('Requiere atención especial y pruebas', style='List Bullet')
    
    doc.add_heading('¿Qué debes hacer?', 3)
    doc.add_paragraph('Asegúrate de que el cambio es intencional', style='List Bullet')
    doc.add_paragraph('Verifica que se probó correctamente', style='List Bullet')
    doc.add_paragraph('Considera hacer backup antes de aplicar', style='List Bullet')
    
    doc.add_heading('📊 Cambio en Modelo', 2)
    doc.add_paragraph('Cambios en la estructura de datos. Por ejemplo, agregar un nuevo campo a una tabla o cambiar cómo se organiza la información.')
    
    doc.add_heading('🌐 Cambio en API', 2)
    doc.add_paragraph('Cambios en cómo otras aplicaciones o servicios se comunican con tu sistema. Una API es como una "interfaz" que permite que diferentes sistemas hablen entre sí.')
    
    doc.add_heading('¿Por qué es importante?', 3)
    doc.add_paragraph('Si cambias una API, otras aplicaciones que la usan podrían romperse', style='List Bullet')
    doc.add_paragraph('Puede afectar integraciones con servicios externos', style='List Bullet')
    doc.add_paragraph('Requiere comunicación con otros equipos si es necesario', style='List Bullet')
    
    doc.add_heading('⚙️ Cambio en Configuración', 2)
    doc.add_paragraph('Cambios en los ajustes del sistema, como URLs, contraseñas, opciones de funcionamiento, etc.')
    
    doc.add_heading('📦 Cambio en Dependencias', 2)
    doc.add_paragraph('Cambios en las "herramientas" o librerías que usa tu código. Son como las herramientas que necesitas para construir algo.')
    
    # SECCIÓN 5: Advertencias
    doc.add_heading('Advertencias y Problemas Potenciales', 1)
    doc.add_paragraph('Esta sección te alerta sobre cosas que deberías revisar antes de guardar tus cambios.')
    
    doc.add_heading('⚠️ Archivo Muy Grande', 2)
    p = doc.add_paragraph('Mensaje ejemplo: ')
    code_run = p.add_run('⚠️ Archivo muy grande: zdash/revi/orders/api/views.py (800 líneas)')
    code_run.font.name = 'Courier New'
    
    doc.add_paragraph('Modificaste un archivo con muchos cambios (más de 500 líneas añadidas o eliminadas en total).')
    
    doc.add_heading('¿Por qué es importante?', 3)
    doc.add_paragraph('Cambios muy grandes son difíciles de revisar', style='List Bullet')
    doc.add_paragraph('Es más probable que contengan errores', style='List Bullet')
    doc.add_paragraph('Son difíciles de entender para otros desarrolladores', style='List Bullet')
    doc.add_paragraph('Si algo sale mal, es difícil identificar el problema', style='List Bullet')
    
    doc.add_heading('¿Qué debes hacer?', 3)
    doc.add_paragraph('Considera dividir el cambio en partes más pequeñas', style='List Bullet')
    doc.add_paragraph('Si es necesario mantenerlo grande, revisa muy cuidadosamente', style='List Bullet')
    doc.add_paragraph('Asegúrate de que el cambio está bien documentado', style='List Bullet')
    
    doc.add_heading('🐛 Código de Debug', 2)
    doc.add_paragraph('El sistema detectó código de "depuración" o "debugging" en tus cambios. Esto incluye cosas como print(), console.log(), debugger, pdb.set_trace().')
    
    doc.add_heading('¿Por qué es importante?', 3)
    doc.add_paragraph('Este código normalmente se usa solo para probar y encontrar errores', style='List Bullet')
    doc.add_paragraph('No debería quedarse en el código final', style='List Bullet')
    doc.add_paragraph('Puede exponer información sensible', style='List Bullet')
    doc.add_paragraph('Puede hacer que el código sea más lento', style='List Bullet')
    
    doc.add_heading('🔒 Posibles Secretos', 2)
    doc.add_paragraph('El sistema detectó palabras que comúnmente se usan para información sensible, como password, api_key, secret, token.')
    
    p = doc.add_paragraph()
    p.add_run('¿Por qué es MUY importante?').bold = True
    
    doc.add_paragraph('NUNCA debes guardar contraseñas, claves o secretos directamente en el código', style='List Bullet')
    doc.add_paragraph('Esto es un riesgo de seguridad grave', style='List Bullet')
    doc.add_paragraph('Si alguien accede al código, puede ver esta información', style='List Bullet')
    doc.add_paragraph('Puede comprometer la seguridad de todo el sistema', style='List Bullet')
    
    doc.add_heading('¿Qué debes hacer?', 3)
    doc.add_paragraph('DETENTE y revisa el archivo inmediatamente', style='List Bullet')
    doc.add_paragraph('Si encontraste un secreto, NO hagas commit hasta arreglarlo', style='List Bullet')
    doc.add_paragraph('Usa variables de entorno o sistemas de gestión de secretos', style='List Bullet')
    doc.add_paragraph('Si accidentalmente ya guardaste un secreto, cámbialo inmediatamente', style='List Bullet')
    
    doc.add_heading('Ejemplo de lo que NO debes hacer:', 3)
    code_para = doc.add_paragraph()
    code_run = code_para.add_run('# ❌ MAL - No hagas esto\nAPI_KEY = "sk_live_1234567890abcdef"\nPASSWORD = "mi_contraseña_secreta"')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    doc.add_heading('Ejemplo de lo que SÍ debes hacer:', 3)
    code_para = doc.add_paragraph()
    code_run = code_para.add_run('# ✅ BIEN - Usa variables de entorno\nAPI_KEY = os.environ.get(\'API_KEY\')\nPASSWORD = os.environ.get(\'DATABASE_PASSWORD\')')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    doc.add_heading('📝 TODOs o FIXMEs', 2)
    doc.add_paragraph('El sistema encontró comentarios que indican trabajo pendiente, como TODO, FIXME, XXX, HACK.')
    
    # SECCIÓN 6: Archivos por Tipo
    doc.add_heading('Archivos por Tipo', 1)
    doc.add_paragraph('Esta sección agrupa tus archivos por categoría para que puedas entender mejor qué tipo de cambios estás haciendo.')
    
    types = [
        ('🐍 Python', '.py', 'Archivos de código Python'),
        ('📜 JavaScript/TypeScript', '.js, .ts, .jsx, .tsx', 'Archivos de código JavaScript o TypeScript'),
        ('🔄 Migración', 'migrations/*.py', 'Archivos que cambian la estructura de la base de datos'),
        ('📊 Modelo', 'models/*.py', 'Archivos que definen estructuras de datos'),
        ('🌐 API', 'api/*.py', 'Archivos que manejan comunicación con otras aplicaciones'),
        ('🧪 Test', 'test*.py, *test.py, tests/', 'Archivos de pruebas automatizadas'),
        ('⚙️ Configuración', 'settings/*.py, config.py, .env', 'Archivos de configuración del sistema'),
        ('📦 Dependencias', 'requirements.txt, pyproject.toml, Pipfile', 'Archivos que listan las herramientas necesarias'),
        ('🐳 Docker', 'Dockerfile, docker-compose.yaml', 'Archivos de configuración de contenedores Docker'),
        ('📚 Documentación', '.md, .rst, .txt', 'Archivos de documentación'),
    ]
    
    for name, ext, desc in types:
        doc.add_heading(name, 2)
        p = doc.add_paragraph()
        p.add_run(f'Extensiones/Patrones: ').bold = True
        p.add_run(ext)
        doc.add_paragraph(desc)
    
    # SECCIÓN 7: Detalles por Archivo
    doc.add_heading('Detalles por Archivo', 1)
    doc.add_paragraph('Esta es la sección más detallada. Muestra cada archivo individual que modificaste con información específica.')
    
    doc.add_heading('Información que Verás para Cada Archivo', 2)
    
    doc.add_heading('1. Nombre del Archivo', 3)
    doc.add_paragraph('La ruta completa del archivo que cambió.')
    
    doc.add_heading('2. Icono de Estado', 3)
    doc.add_paragraph('Un icono de color que indica qué tipo de cambio hiciste:')
    doc.add_paragraph('➕ Verde (A) - Añadido: Archivo NUEVO que agregaste al proyecto', style='List Bullet')
    doc.add_paragraph('✏️ Azul (M) - Modificado: Archivo que ya existía y que CAMBIASTE', style='List Bullet')
    doc.add_paragraph('🗑️ Rojo (D) - Eliminado: Archivo que ELIMINASTE del proyecto', style='List Bullet')
    doc.add_paragraph('📝 Amarillo (R) - Renombrado: Archivo que RENOMBRASTE o moviste a otra ubicación', style='List Bullet')
    
    doc.add_heading('3. Etiqueta de Tipo', 3)
    doc.add_paragraph('Una pequeña etiqueta que dice qué tipo de archivo es (python, api, model, test, config, etc.)')
    
    doc.add_heading('4. Etiqueta de Estado', 3)
    doc.add_paragraph('Una etiqueta que dice el estado del cambio (Añadido, Modificado, Eliminado, Renombrado)')
    
    doc.add_heading('5. Números de Cambio', 3)
    doc.add_paragraph('Dos números que muestran las líneas que cambiaron:')
    doc.add_paragraph('Verde (+X): Líneas que AGREGASTE en ese archivo', style='List Bullet')
    doc.add_paragraph('Rojo (-X): Líneas que ELIMINASTE en ese archivo', style='List Bullet')
    
    doc.add_paragraph('Ejemplo: +45 / -12 significa que agregaste 45 líneas nuevas y eliminaste 12 líneas viejas. El archivo creció en 33 líneas netas.')
    
    doc.add_heading('6. Botón "Ver diff"', 3)
    doc.add_paragraph('Un botón que te permite ver las líneas exactas que cambiaste.')
    doc.add_paragraph('Al hacer clic, verás:', style='List Bullet')
    doc.add_paragraph('Las líneas en VERDE con un + al inicio: Son las líneas que AGREGASTE', style='List Bullet 2')
    doc.add_paragraph('Las líneas en ROJO con un - al inicio: Son las líneas que ELIMINASTE', style='List Bullet 2')
    doc.add_paragraph('Las líneas sin color: Son líneas de contexto que no cambiaron', style='List Bullet 2')
    
    # SECCIÓN 8: Crear Commit
    doc.add_heading('Crear un Commit', 1)
    doc.add_paragraph('Esta es la sección final donde guardas tus cambios permanentemente en el historial del proyecto.')
    
    doc.add_heading('Campo de Texto: "Mensaje del commit"', 2)
    doc.add_paragraph('Un mensaje corto (idealmente 50-72 caracteres) que describe qué hiciste y por qué (si es relevante).')
    
    doc.add_heading('¿Por qué es importante?', 3)
    doc.add_paragraph('Este mensaje queda guardado para siempre en el historial', style='List Bullet')
    doc.add_paragraph('Ayuda a otros desarrolladores (y a ti en el futuro) a entender qué cambió', style='List Bullet')
    doc.add_paragraph('Facilita encontrar cambios específicos más adelante', style='List Bullet')
    doc.add_paragraph('Es parte de la documentación del proyecto', style='List Bullet')
    
    doc.add_heading('Buenos ejemplos de mensajes:', 3)
    good_examples = [
        'Agregar funcionalidad de búsqueda de usuarios',
        'Corregir error al calcular totales de pedidos',
        'Actualizar dependencias de seguridad',
        'Mejorar rendimiento de consultas a base de datos',
        'Agregar validación de email en formulario de registro'
    ]
    for ex in good_examples:
        p = doc.add_paragraph(ex, style='List Bullet')
        p.style.font.color.rgb = RGBColor(0, 128, 0)  # Verde
    
    doc.add_heading('Malos ejemplos de mensajes:', 3)
    bad_examples = ['cambios', 'fix', 'update', 'asdf', 'wip']
    for ex in bad_examples:
        p = doc.add_paragraph(ex, style='List Bullet')
        p.style.font.color.rgb = RGBColor(255, 0, 0)  # Rojo
    
    doc.add_heading('Consejos para escribir buenos mensajes:', 3)
    tips = [
        'Sé específico: Di qué cambió, no solo "cambios"',
        'Usa el imperativo: "Agregar" en lugar de "Agregué" o "Se agregó"',
        'Sé conciso: Idealmente una línea, máximo 72 caracteres',
        'Explica el "por qué" si es necesario: A veces ayuda saber la razón del cambio'
    ]
    for tip in tips:
        doc.add_paragraph(tip, style='List Bullet')
    
    # SECCIÓN 9: Ejemplos Prácticos
    doc.add_heading('Ejemplos Prácticos', 1)
    
    doc.add_heading('Ejemplo 1: Agregar una Nueva Funcionalidad', 2)
    doc.add_paragraph('Escenario: Agregaste una nueva funcionalidad de búsqueda de usuarios.')
    
    doc.add_heading('Reporte que verías:', 3)
    report_text = """📊 ESTADÍSTICAS:
- Archivos: 8 archivos
- Añadidas: +320 líneas
- Eliminadas: -45 líneas
- Cambio Neto: +275 líneas

✨ CAMBIOS SIGNIFICATIVOS:
- 🌐 Cambio en API: zdash/revi/user/api/views.py
- 📊 Cambio en modelo: zdash/revi/user/models.py

📁 ARCHIVOS POR TIPO:
- python: 6 archivos
- api: 2 archivos
- test: 2 archivos"""
    
    code_para = doc.add_paragraph(report_text)
    code_run = code_para.runs[0]
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    doc.add_heading('Interpretación:', 3)
    doc.add_paragraph('Cambio de tamaño medio (8 archivos, 275 líneas netas)', style='List Bullet')
    doc.add_paragraph('Agregaste pruebas (2 archivos de test - buena práctica)', style='List Bullet')
    doc.add_paragraph('Cambios en API y modelo (esperado para nueva funcionalidad)', style='List Bullet')
    doc.add_paragraph('No hay advertencias (todo se ve bien)', style='List Bullet')
    doc.add_paragraph('El cambio neto es positivo (agregaste funcionalidad)', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Mensaje de commit sugerido: ').bold = True
    p.add_run('Agregar funcionalidad de búsqueda de usuarios')
    
    # Más ejemplos...
    doc.add_heading('Ejemplo 2: Corrección de un Error', 2)
    doc.add_paragraph('Escenario: Corregiste un error en el cálculo de totales.')
    
    doc.add_heading('Ejemplo 3: Cambio Grande con Advertencias', 2)
    doc.add_paragraph('Escenario: Refactorizaste una parte grande del código.')
    
    doc.add_heading('Ejemplo 4: Actualización de Dependencias', 2)
    doc.add_paragraph('Escenario: Actualizaste las herramientas que usa el proyecto.')
    
    doc.add_heading('Ejemplo 5: Problema de Seguridad Detectado', 2)
    doc.add_paragraph('Escenario: El sistema detectó posible información sensible.')
    doc.add_paragraph('🚨 ALERTA DE SEGURIDAD: Posible secreto detectado')
    doc.add_paragraph('🚨 NO HAGAS COMMIT hasta revisar esto')
    
    # SECCIÓN 10: Preguntas Frecuentes
    doc.add_heading('Preguntas Frecuentes', 1)
    
    faqs = [
        ('¿Qué hago si veo una advertencia de "secretos"?', 
         'DETENTE inmediatamente. Revisa el archivo y asegúrate de que no estés guardando contraseñas, claves API, o cualquier información sensible directamente en el código. Si necesitas guardar configuración sensible, usa variables de entorno o sistemas de gestión de secretos. NUNCA hagas commit con secretos en el código.'),
        ('¿Es malo tener muchos archivos modificados?', 
         'No necesariamente, pero puede ser difícil de revisar. Si son más de 20 archivos, considera si puedes dividir este cambio en partes más pequeñas, si todos los archivos están relacionados con el mismo cambio, y si es necesario mantenerlo grande, asegúrate de revisar cuidadosamente cada archivo.'),
        ('¿Qué significa "Cambio Neto" negativo?', 
         'Significa que eliminaste más código del que agregaste. Esto puede ser bueno (si estás limpiando código viejo, eliminando código duplicado, o refactorizando) o requiere revisión (asegúrate de que no eliminaste funcionalidad que aún se necesita).'),
        ('¿Debo revisar todos los diffs (cambios línea por línea)?', 
         'Idealmente sí, especialmente si son cambios importantes. Pero al menos revisa los archivos con advertencias, los archivos que tienen muchos cambios (más de 100 líneas), los archivos de cambios significativos (APIs, modelos, configuración), y una muestra de los archivos más pequeños para asegurarte de que todo está bien.'),
        ('¿Qué pasa si no hay archivos en staging?', 
         'Necesitas primero seleccionar qué archivos quieres guardar. Puedes usar el botón "Stage All" en la interfaz para añadir todos los archivos modificados, o usar git add en la terminal para añadir archivos específicos.'),
        ('¿Puedo hacer commit sin escribir un mensaje?', 
         'Técnicamente algunos sistemas lo permiten, pero NO es recomendable. El mensaje del commit es importante porque ayuda a otros (y a ti en el futuro) a entender qué cambió, facilita encontrar cambios específicos, y es parte de la documentación del proyecto.'),
    ]
    
    for question, answer in faqs:
        doc.add_heading(question, 2)
        doc.add_paragraph(answer)
    
    # SECCIÓN 11: Glosario
    doc.add_heading('Glosario de Términos', 1)
    
    glossary = [
        ('API (Application Programming Interface)', 
         'Una interfaz que permite que diferentes aplicaciones o servicios se comuniquen entre sí. Es como un "contrato" que define cómo pueden interactuar dos sistemas.'),
        ('Branch (Rama)', 
         'Una versión separada del código donde puedes trabajar sin afectar la versión principal. Es como tener diferentes "carpetas" del mismo proyecto.'),
        ('Commit', 
         'El acto de guardar un conjunto de cambios en el historial del proyecto. Es como hacer un "guardado" permanente con un mensaje que describe qué cambió.'),
        ('Diff', 
         'La diferencia entre la versión antigua y nueva de un archivo. Muestra exactamente qué líneas cambiaron, cuáles se agregaron, y cuáles se eliminaron.'),
        ('Dependencias', 
         'Las "herramientas" o librerías que tu código necesita para funcionar. Son como las herramientas que necesitas para construir algo.'),
        ('Migración', 
         'Un cambio en la estructura de la base de datos. Define cómo se organiza y almacena la información.'),
        ('Modelo', 
         'Una definición de cómo se estructura un tipo de dato en el sistema. Por ejemplo, un "modelo de Usuario" define qué información se guarda sobre cada usuario.'),
        ('Staging', 
         'El proceso de preparar archivos para ser guardados. Es como poner archivos en una "bandeja" antes de guardarlos permanentemente.'),
        ('Variable de Entorno', 
         'Una forma segura de guardar información sensible (como contraseñas) fuera del código. El código lee estas variables cuando se ejecuta, pero no están guardadas en el código mismo.'),
    ]
    
    for term, definition in glossary:
        p = doc.add_paragraph()
        p.add_run(f'{term}: ').bold = True
        p.add_run(definition)
    
    # CONCLUSIÓN
    doc.add_heading('Conclusión', 1)
    doc.add_paragraph('Smart Commit es una herramienta poderosa que te ayuda a:')
    doc.add_paragraph('Entender exactamente qué estás guardando', style='List Bullet')
    doc.add_paragraph('Detectar posibles problemas antes de que ocurran', style='List Bullet')
    doc.add_paragraph('Mantener un historial claro y documentado', style='List Bullet')
    doc.add_paragraph('Trabajar con más confianza y seguridad', style='List Bullet')
    
    doc.add_heading('Recuerda:', 2)
    doc.add_paragraph('Siempre revisa las advertencias, especialmente las de seguridad', style='List Bullet')
    doc.add_paragraph('Escribe mensajes de commit claros y descriptivos', style='List Bullet')
    doc.add_paragraph('Revisa los cambios grandes cuidadosamente', style='List Bullet')
    doc.add_paragraph('Cuando tengas dudas, pregunta a un desarrollador', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('¡Feliz desarrollo! 🚀').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Pie de página
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('Este documento fue creado para ayudar a usuarios no técnicos a entender y usar Smart Commit efectivamente.')
    set_run_font(footer_run, font_size=9, color=RGBColor(128, 128, 128))
    
    return doc


def main():
    """Función principal"""
    print("📝 Generando documento Word...")
    
    try:
        doc = create_document()
        output_file = "Smart_Commit_Guia_Completa.docx"
        doc.save(output_file)
        print(f"✅ Documento creado exitosamente: {output_file}")
        print(f"📄 Ubicación: {output_file}")
        return 0
    except ImportError:
        print("❌ Error: Necesitas instalar python-docx")
        print("   Ejecuta: pip install python-docx")
        return 1
    except Exception as e:
        print(f"❌ Error al generar el documento: {e}")
        return 1


if __name__ == '__main__':
    sys.exit(main())
